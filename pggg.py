import streamlit as st
from eventregistry import EventRegistry, QueryArticlesIter, QueryItems
import json
from docx import Document
from datetime import datetime, timedelta
import io
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Initialization block
keywords_list = ["Zelensky"]
# keywords_list = ["Zelensky","Human Rights", "Illegal occupied", "Anti-India", "Hindu", "Supremacy", "Nationalist", "India", "USA", "China", "Bangladesh", "UK", "Lanka", "Afghanistan", "Militant", "Khalistan", "terror"]
journal_list = ["washingtonpost.com"]
language_list = ["eng", "hin"]
topics = ["Geo Politics", "Military", "Development", "Health", "Business"]
ignore_topic_list = ['hockey', 'cricket', "Bollywood", "Hollywood", "Box office", "Asia Cup", "Games", "sport", "Fashion"]
author_names = [
   "Lizzie Johnson", "Serhiy Morgunov"]
#   , "Kirti Dubey", "John Reed", "Mallika Sen", "Ronojoy Mazumdar", "Chiranjivi Chakraborty", "Sibi Arasu", "Jere Longman", "Suhasini Raj", "Siddhartha Singh", 
 #   "Chris Kay", "Sankalp Phartiyal", "Shruti Srivastava", "Ashutosh Joshi", "Ng Wei Kai", "Anup Sinha", "Sudhi Ranjan Sen", 
 #   "Ruchi Bhatia", "Peter Martin", "Iain Marlow", "Martin Wolf", "Alex Gabriel Simon", "Amrit Dhillon", "Mihir Sharma", "Anumita Kaur"
# ]
max_items_per_journal = 5  # Maximum 5 articles per newspaper
max_total_articles = 10  # Maximum total articles to retrieve
days_range = 1
date_end = datetime.now().strftime('%Y-%m-%d')
date_start = (datetime.now() - timedelta(days=days_range)).strftime('%Y-%m-%d')

# Initialize a set to store unique article URLs
unique_article_urls = set()

# Function to determine if an article should be ignored
def should_ignore_article(article):
    title = article.get("title", "").lower()
    body = article.get("body", "").lower()
    ignore_keywords = ["World Cup", "cricket", "football", "tennis", "match", "Asia cup", "Kohli", "apple", "Asia Cup", "Top News", "badminton"]
    
    for keyword in ignore_keywords:
        if keyword.lower() in title or keyword.lower() in body:
            return True
    return False

# Function to retrieve articles for a given keyword and journal
def retrieve_articles(keyword, journal):
    try:
        er = EventRegistry(apiKey=st.secrets["secret_key"], allowUseOfArchive=False)
        q = QueryArticlesIter(
            keywords=QueryItems.AND([keyword]),  # Use AND for the main keyword
            sourceUri=er.getSourceUri(journal),
            lang=QueryItems.OR(language_list),
            ignoreKeywords=QueryItems.OR(ignore_topic_list),
            authorUri=QueryItems.OR(author_names),  # Add author names filter here
            isDuplicateFilter="skipDuplicates",
            dataType="news",
            dateStart=date_start,
            dateEnd=date_end
        )

        print(f"Number of results for '{keyword}' in '{journal}': {q.count(er)}")

        # Store the articles in the 'articles' list
        articles = []
        count = 0
        for art in q.execQuery(er, sortBy="rel"):
            articles.append(json.dumps(art, indent=4))
            count += 1
            if count >= max_items_per_journal or count >= max_total_articles:
                break  # Stop after reaching the max articles limit

        return articles
    except Exception as e:
        print(f"An error occurred for '{keyword}' in '{journal}': {e}")
        return []

def process_keyword_journal_combination(args):
    keyword, journal = args
    articles = retrieve_articles(keyword, journal)
    return articles

def main():
    st.title("Digital Press Clipping Generator")

    # Initialize all_articles as an empty list
    all_articles = []
    doc = Document()

    # Add a button to trigger document generation
    if st.button("Generate Digital Press Clipping"):
        doc, all_articles = generate_document()
        st.success("Document generated successfully!")

    # Display the document content
    if all_articles:
        st.subheader("Generated Document Content:")
        for article in all_articles:
            art = json.loads(article)
            title = art["title"]
            source = art["url"]
            content = art["body"]

            st.markdown(f"**Title:** {title}")
            st.markdown(f"**Source:** {source}")
            st.markdown(f"**Content:** {content}")
            st.markdown("---")

    # Add a download button for the generated document
    bio = io.BytesIO()
    doc.save(bio)
    if doc:
        st.download_button(
            label="Click here to download the generated document",
            data=bio.getvalue(),
            file_name="DigitalClippings.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

def generate_document():
    # Fetch and add articles to the document
    doc = Document()  # Initialize the document
    all_articles = []
    total_articles_count = 0

    # Add the front page
    front_page = doc.add_paragraph()
    front_page.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center align the text

    title = front_page.add_run("DIGITAL MEDIA PRESS CLIPPING")
    title.bold = True
    title.font.size = Pt(18)

    # Get the current date and day
    current_date = datetime.now().strftime('%Y-%m-%d')
    current_day = datetime.now().strftime('%A')

    # Add the date and day on the next line
    date_and_day = front_page.add_run(f"\n{current_date} ({current_day})")
    date_and_day.font.size = Pt(14)

    # Add a page break after the front page
    doc.add_page_break()

    # Continue with adding articles
    for keyword in keywords_list:
        for journal in journal_list:
            articles = retrieve_articles(keyword, journal)
            for article in articles:
                art = json.loads(article)
                article_url = art.get("url")
                if article_url not in unique_article_urls:
                    all_articles.append(article)
                    unique_article_urls.add(article_url)

            total_articles_count += len(articles)
            if total_articles_count >= max_total_articles:
                break  # Stop if the maximum total articles limit is reached

    for article in all_articles:
        art = json.loads(article)
        title = art["title"]
        source = art["url"]
        content = art["body"]

        # Add article information to the document
        doc.add_heading(title, level=2)
        doc.add_paragraph(f"Source: {source}")
        doc.add_paragraph("Content:").alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        doc.add_paragraph(content).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        doc.add_paragraph("==================")

    return doc, all_articles

if __name__ == "__main__":
    main()
